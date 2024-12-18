import streamlit as st
import openai
from docx import Document
from io import BytesIO

# OpenAI API key setup (use secrets or environment variable for security)
openai.api_key = st.secrets["openai"]["api_key"]

# Function to generate PIP using OpenAI API (Updated for version 0.28+)
def generate_pip(difficulty, explanation):
    prompt = f"""
    The user is a medical student in a pediatric clerkship and is experiencing difficulty in the area of {difficulty}.
    They mentioned the following explanation for their difficulty: {explanation if explanation else 'No additional explanation provided.'}
    
    Generate a detailed Performance Improvement Plan (PIP) that includes:
    - A summary of the identified issue
    - Specific strategies to improve in this area
    - Recommended resources or actions for the student
    """

    # Call OpenAI API using the ChatCompletion method (correct for version 0.28+)
    response = openai.ChatCompletion.create(
        model="gpt-4",  # Or use "gpt-3.5-turbo" depending on your preference
        messages=[
            {"role": "system", "content": "You are an expert in pediatric medical education."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1000,  # Adjust based on your preference
    )

    # Extract the PIP text from the response
    pip_text = response['choices'][0]['message']['content'].strip()
    return pip_text

# Function to create a Word document from the PIP
def create_word_doc(pip_text):
    doc = Document()
    doc.add_heading('Performance Improvement Plan', 0)
    doc.add_paragraph(pip_text)
    
    # Save to a BytesIO buffer to be sent to the user
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# Streamlit app
def app():
    st.title("Pediatric Clerkship Reflection Application")
    st.subheader("Growing through experience, excelling through practice")
    
    st.markdown("""
    Welcome to the Pediatric Clerkship Reflection Application. Please reflect on your current challenges and identify areas where you're experiencing difficulty. Based on your inputs, we will help generate a Performance Improvement Plan (PIP) to guide your growth.
    """)

    # Dropdown to select the area of difficulty
    difficulty_options = [
        "Knowledge for Practice", 
        "Clinical Reasoning", 
        "Oral Presentation", 
        "Documentation", 
        "Communication with Families", 
        "Communication with Team Members",
        "Time Management",
        "Patient Interaction",
        "Clinical Skills",
        "Professionalism",
        "Teamwork",
        "Stress Management",
        "Decision Making"
    ]
    difficulty = st.selectbox("Select the area where you are experiencing difficulties:", difficulty_options + ["Other"])
    
    # Text input for explanation (optional)
    explanation = st.text_area("If applicable, why do you think you are experiencing difficulties?", "")

    # Button to generate the PIP
    if st.button("Push PIP Button"):
        if difficulty == "Other":
            difficulty = st.text_input("Please specify the area where you're experiencing difficulties:")

        if not difficulty:
            st.error("Please select or specify an area of difficulty.")
        else:
            st.info("Generating your Performance Improvement Plan... Please wait.")
            
            # Generate the PIP text
            pip_text = generate_pip(difficulty, explanation)
            
            # Create a Word document
            word_doc = create_word_doc(pip_text)
            
            # Provide the Word document as a download
            st.download_button(
                label="Download your PIP document",
                data=word_doc,
                file_name="Performance_Improvement_Plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            st.success("Your PIP has been successfully created!")

# Run the app
if __name__ == "__main__":
    app()
